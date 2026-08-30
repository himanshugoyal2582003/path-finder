import docx

doc = docx.Document("PathFinder_—_Product_&_Visual_Design_Specification.docx")

with open("PathFinder_design_spec.md", "w", encoding="utf-8") as f:
    for para in doc.paragraphs:
        if para.text.strip():
            f.write(para.text + "\n\n")
    
    # Let's also extract tables if there are any
    for i, table in enumerate(doc.tables):
        f.write(f"\n### Table {i+1}\n\n")
        for row in table.rows:
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            f.write("| " + " | ".join(row_text) + " |\n")
        f.write("\n")

print("Successfully extracted docx contents to PathFinder_design_spec.md")
